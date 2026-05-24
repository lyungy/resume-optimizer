"""
DOCX 文件生成工具 - 精美模板版
简历：专业简洁风，ATS友好
面试攻略：结构清晰，视觉层次分明
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime


# ============================================================
# 配色方案
# ============================================================
class Colors:
    """专业配色"""
    PRIMARY = RGBColor(0x2B, 0x57, 0x9A)       # 主色 - 深蓝
    PRIMARY_LIGHT = RGBColor(0x3A, 0x7C, 0xBD)  # 主色浅
    ACCENT = RGBColor(0x1A, 0x8C, 0x6B)         # 强调色 - 墨绿
    ACCENT_LIGHT = RGBColor(0xE8, 0xF5, 0xE9)   # 强调色底
    TEXT_PRIMARY = RGBColor(0x2D, 0x2D, 0x2D)    # 主文字 - 近黑
    TEXT_SECONDARY = RGBColor(0x66, 0x66, 0x66)  # 次文字 - 灰
    TEXT_LIGHT = RGBColor(0x99, 0x99, 0x99)      # 浅文字
    DIVIDER = RGBColor(0xDD, 0xDD, 0xDD)         # 分割线
    BG_LIGHT = RGBColor(0xF5, 0xF7, 0xFA)        # 浅底色
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    HIGH = RGBColor(0xE8, 0x4D, 0x3D)            # 高优先级 - 红
    MEDIUM = RGBColor(0xF5, 0xA6, 0x23)          # 中优先级 - 橙
    LOW = RGBColor(0x4C, 0xAF, 0x50)             # 低优先级 - 绿


# ============================================================
# 通用工具函数
# ============================================================
def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_bottom_border(paragraph, color: str = "2B579A", width: int = 6):
    """给段落添加下边框"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_top_border(paragraph, color: str = "2B579A", width: int = 4):
    """给段落添加上边框"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{width}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: float = 1.15):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _add_run(paragraph, text: str, bold: bool = False, size: int = 10,
             color: RGBColor = None, font_name: str = None):
    """添加格式化文本"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def _set_run_font(run, name: str = "微软雅黑", size: int = 10,
                  bold: bool = False, color: RGBColor = None):
    """设置 run 字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _add_section_header(doc, title: str, color: str = "2B579A"):
    """添加带装饰的章节标题"""
    # 空行
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=0)

    # 标题行
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=0, after=4)
    _add_bottom_border(p, color, width=8)
    run = p.add_run(f"  {title}")
    _set_run_font(run, "微软雅黑", 13, bold=True, color=Colors.PRIMARY)

    return p


def _add_subsection(doc, title: str):
    """添加子标题"""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=2)
    run = p.add_run(f"▸ {title}")
    _set_run_font(run, "微软雅黑", 11, bold=True, color=Colors.TEXT_PRIMARY)
    return p


def _add_bullet(doc, text: str, indent_level: int = 0):
    """添加美化项目符号"""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=1, after=1, line=1.3)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3 + indent_level * 0.2)
    pf.first_line_indent = Inches(-0.15)

    run = p.add_run("● ")
    _set_run_font(run, "微软雅黑", 7, color=Colors.PRIMARY_LIGHT)

    run = p.add_run(text)
    _set_run_font(run, "微软雅黑", 10, color=Colors.TEXT_PRIMARY)
    return p


def _add_tag_paragraph(doc, tags: list[str]):
    """添加标签式展示（技能标签）"""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=2, after=4)

    for i, tag in enumerate(tags):
        if i > 0:
            run = p.add_run("  ·  ")
            _set_run_font(run, "微软雅黑", 9, color=Colors.DIVIDER)
        run = p.add_run(tag)
        _set_run_font(run, "微软雅黑", 9, color=Colors.PRIMARY)
    return p


# ============================================================
# 简历生成
# ============================================================
class DocxWriter:
    """DOCX 文件生成器"""

    def create_resume(
        self,
        content: dict,
        output_path: str,
        original_doc_path: str = None,
    ) -> str:
        """
        生成优化后的精美简历
        """
        doc = Document()

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(10)
        style.font.color.rgb = Colors.TEXT_PRIMARY
        style.paragraph_format.line_spacing = 1.15

        # ---- 姓名标题 ----
        name = content.get("name", "")
        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=12, after=2)
            run = p.add_run(name)
            _set_run_font(run, "微软雅黑", 22, bold=True, color=Colors.PRIMARY)

            # 装饰线
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=0, after=8)
            run = p.add_run("━" * 30)
            _set_run_font(run, "微软雅黑", 8, color=Colors.DIVIDER)

        # ---- 匹配度评分 ----
        match_score = content.get("match_score")
        if match_score:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=0, after=6)

            score_color = Colors.ACCENT if match_score >= 70 else Colors.MEDIUM if match_score >= 50 else Colors.HIGH
            run = p.add_run(f"JD 匹配度：{match_score}分")
            _set_run_font(run, "微软雅黑", 10, bold=True, color=score_color)

        optimized = content.get("optimized_sections", {})

        # ---- 个人总结 ----
        summary = optimized.get("summary", "")
        if summary:
            _add_section_header(doc, "个人总结")
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=4, line=1.4)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.1)
            run = p.add_run(summary)
            _set_run_font(run, "微软雅黑", 10, color=Colors.TEXT_PRIMARY)

        # ---- 专业技能 ----
        skills = optimized.get("skills", [])
        if skills:
            _add_section_header(doc, "专业技能")
            _add_tag_paragraph(doc, skills)

        # ---- 工作经历 ----
        experience = optimized.get("experience", [])
        if experience:
            _add_section_header(doc, "工作经历")
            for exp in experience:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=6, after=2)

                # 公司名
                company = exp.get("company", "")
                title = exp.get("title", "")
                period = exp.get("period", "")

                run = p.add_run(company)
                _set_run_font(run, "微软雅黑", 11, bold=True, color=Colors.TEXT_PRIMARY)

                if title:
                    run = p.add_run(f"  |  {title}")
                    _set_run_font(run, "微软雅黑", 10, color=Colors.TEXT_SECONDARY)

                if period:
                    run = p.add_run(f"  ({period})")
                    _set_run_font(run, "微软雅黑", 9, color=Colors.TEXT_LIGHT)

                # 亮点
                for highlight in exp.get("highlights", []):
                    _add_bullet(doc, highlight)

        # ---- 项目经历 ----
        projects = optimized.get("projects", [])
        if projects:
            _add_section_header(doc, "项目经历")
            for proj in projects:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=6, after=2)

                name = proj.get("name", "")
                role = proj.get("role", "")

                run = p.add_run(name)
                _set_run_font(run, "微软雅黑", 11, bold=True, color=Colors.TEXT_PRIMARY)

                if role:
                    run = p.add_run(f"  |  {role}")
                    _set_run_font(run, "微软雅黑", 10, color=Colors.TEXT_SECONDARY)

                for highlight in proj.get("highlights", []):
                    _add_bullet(doc, highlight)

        # ---- 优化建议 ----
        suggestions = content.get("suggestions", [])
        if suggestions:
            _add_section_header(doc, "优化建议", color="1A8C6B")
            for suggestion in suggestions:
                _add_bullet(doc, suggestion)

        # ---- ATS 建议 ----
        ats_tips = content.get("ats_tips", [])
        if ats_tips:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=2)
            run = p.add_run("ATS 通过率提升建议")
            _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.ACCENT)

            for tip in ats_tips:
                _add_bullet(doc, tip)

        # ---- 页脚 ----
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=12, after=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"由简历优化系统生成 · {datetime.now().strftime('%Y-%m-%d')}")
        _set_run_font(run, "微软雅黑", 8, color=Colors.TEXT_LIGHT)

        # 保存
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        return str(output)

    # ============================================================
    # 面试攻略生成
    # ============================================================
    def create_interview_guide(
        self,
        content: dict,
        output_path: str,
        jd_title: str = "",
        company_name: str = "",
    ) -> str:
        """
        生成精美的面试攻略
        """
        doc = Document()

        # 页面设置
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 默认字体
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(10)
        style.font.color.rgb = Colors.TEXT_PRIMARY

        # ---- 封面标题 ----
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=40, after=6)
        run = p.add_run("面 试 攻 略")
        _set_run_font(run, "微软雅黑", 28, bold=True, color=Colors.PRIMARY)

        if company_name or jd_title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, before=4, after=2)
            subtitle = f"{company_name}  ·  {jd_title}" if company_name and jd_title else company_name or jd_title
            run = p.add_run(subtitle)
            _set_run_font(run, "微软雅黑", 14, color=Colors.TEXT_SECONDARY)

        # 装饰线
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=8, after=4)
        run = p.add_run("━" * 40)
        _set_run_font(run, "微软雅黑", 8, color=Colors.DIVIDER)

        # 生成时间
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=2, after=20)
        run = p.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        _set_run_font(run, "微软雅黑", 9, color=Colors.TEXT_LIGHT)

        # ---- 知识点清单 ----
        knowledge_points = content.get("knowledge_points", [])
        if knowledge_points:
            _add_section_header(doc, "📚 知识点清单")

            for kp in knowledge_points:
                category = kp.get("category", "未分类")
                priority = kp.get("priority", "medium")
                hours = kp.get("estimated_prep_hours", 0)

                # 优先级颜色
                pri_color = {"high": Colors.HIGH, "medium": Colors.MEDIUM, "low": Colors.LOW}.get(priority, Colors.MEDIUM)
                pri_label = {"high": "高", "medium": "中", "low": "低"}.get(priority, "中")

                _add_subsection(doc, category)

                # 优先级和时间
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=2)
                pf = p.paragraph_format
                pf.left_indent = Inches(0.2)

                run = p.add_run(f"优先级：{pri_label}  ")
                _set_run_font(run, "微软雅黑", 9, bold=True, color=pri_color)

                run = p.add_run(f"准备时间：{hours}小时")
                _set_run_font(run, "微软雅黑", 9, color=Colors.TEXT_SECONDARY)

                # 知识点列表
                for point in kp.get("points", []):
                    _add_bullet(doc, point, indent_level=1)

                # 学习资源
                resources = kp.get("study_resources", [])
                if resources:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=2, after=1)
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.2)
                    run = p.add_run("📖 学习方向：")
                    _set_run_font(run, "微软雅黑", 9, bold=True, color=Colors.ACCENT)

                    for resource in resources:
                        _add_bullet(doc, resource, indent_level=1)

        # ---- 高频面试题 ----
        questions = content.get("high_frequency_questions", [])
        if questions:
            _add_section_header(doc, "🎯 高频面试题")

            for i, q in enumerate(questions, 1):
                question = q.get("question", "")
                category = q.get("category", "")
                difficulty = q.get("difficulty", "")

                # 题目
                _add_subsection(doc, f"Q{i}: {question}")

                # 元信息
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=2)
                pf = p.paragraph_format
                pf.left_indent = Inches(0.2)

                diff_color = {"hard": Colors.HIGH, "medium": Colors.MEDIUM, "easy": Colors.LOW}.get(difficulty, Colors.MEDIUM)
                diff_label = {"hard": "难", "medium": "中", "easy": "易"}.get(difficulty, "中")

                run = p.add_run(f"分类：{category}  ")
                _set_run_font(run, "微软雅黑", 9, color=Colors.TEXT_SECONDARY)
                run = p.add_run(f"难度：{diff_label}")
                _set_run_font(run, "微软雅黑", 9, bold=True, color=diff_color)

                # 答题模板
                answer = q.get("answer_template", "")
                if answer:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=3, after=1)
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.2)
                    run = p.add_run("💡 答题思路：")
                    _set_run_font(run, "微软雅黑", 9, bold=True, color=Colors.PRIMARY)

                    # 答案内容 - 用浅色背景表格模拟卡片
                    table = doc.add_table(rows=1, cols=1)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell = table.cell(0, 0)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(answer)
                    _set_run_font(run, "微软雅黑", 9.5, color=Colors.TEXT_PRIMARY)
                    _set_cell_shading(cell, "F5F7FA")

                    # 设置表格宽度
                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="9000" w:type="dxa"/>')
                    tblPr.append(tblW)

                # 回答要点
                key_points = q.get("key_points", [])
                if key_points:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=3, after=1)
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.2)
                    run = p.add_run("✅ 回答要点：")
                    _set_run_font(run, "微软雅黑", 9, bold=True, color=Colors.ACCENT)

                    for point in key_points:
                        _add_bullet(doc, point, indent_level=1)

                # 常见错误
                mistakes = q.get("common_mistakes", [])
                if mistakes:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, before=3, after=1)
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.2)
                    run = p.add_run("⚠️ 常见错误：")
                    _set_run_font(run, "微软雅黑", 9, bold=True, color=Colors.HIGH)

                    for mistake in mistakes:
                        _add_bullet(doc, mistake, indent_level=1)

        # ---- 准备策略 ----
        strategy = content.get("preparation_strategy", {})
        if strategy:
            _add_section_header(doc, "📅 准备策略", color="1A8C6B")

            total_days = strategy.get("total_days", 7)
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=4)
            run = p.add_run(f"建议准备周期：{total_days} 天")
            _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.TEXT_PRIMARY)

            # 每日计划 - 用表格展示
            daily_plan = strategy.get("daily_plan", [])
            if daily_plan:
                _add_subsection(doc, "每日计划")

                table = doc.add_table(rows=1, cols=4)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # 表头
                headers = ["天数", "主题", "时长", "任务"]
                header_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    header_cells[i].text = ""
                    p = header_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(h)
                    _set_run_font(run, "微软雅黑", 9, bold=True, color=Colors.WHITE)
                    _set_cell_shading(header_cells[i], "2B579A")

                # 数据行
                for day in daily_plan:
                    row = table.add_row()
                    cells = row.cells

                    day_num = str(day.get("day", ""))
                    focus = day.get("focus", "")
                    hours = f"{day.get('hours', 0)}h"
                    tasks = "、".join(day.get("tasks", []))

                    for j, val in enumerate([day_num, focus, hours, tasks]):
                        cells[j].text = ""
                        p = cells[j].paragraphs[0]
                        run = p.add_run(val)
                        _set_run_font(run, "微软雅黑", 9, color=Colors.TEXT_PRIMARY)

                # 设置表格样式
                self._style_table(table)

            # 面试技巧
            tips = strategy.get("tips", [])
            if tips:
                _add_subsection(doc, "面试技巧")
                for tip in tips:
                    _add_bullet(doc, tip)

        # ---- 公司调研 ----
        company_research = content.get("company_research", {})
        if company_research:
            _add_section_header(doc, "🏢 公司调研")

            what_to_prepare = company_research.get("what_to_prepare", [])
            if what_to_prepare:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=2)
                run = p.add_run("需要了解的信息：")
                _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.TEXT_PRIMARY)

                for item in what_to_prepare:
                    _add_bullet(doc, item)

            questions_to_ask = company_research.get("questions_to_ask", [])
            if questions_to_ask:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=4, after=2)
                run = p.add_run("🎤 可以反问面试官：")
                _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.PRIMARY)

                for q in questions_to_ask:
                    _add_bullet(doc, q)

            red_flags = company_research.get("red_flags", [])
            if red_flags:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=4, after=2)
                run = p.add_run("🚩 注意信号：")
                _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.HIGH)

                for flag in red_flags:
                    _add_bullet(doc, flag)

        # ---- 薪资谈判 ----
        salary = content.get("salary_negotiation", {})
        if salary:
            _add_section_header(doc, "💰 薪资谈判参考")

            market_range = salary.get("market_range", "")
            if market_range:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=2)
                run = p.add_run(f"市场薪资范围：{market_range}")
                _set_run_font(run, "微软雅黑", 10, bold=True, color=Colors.ACCENT)

            neg_tips = salary.get("negotiation_tips", [])
            if neg_tips:
                for tip in neg_tips:
                    _add_bullet(doc, tip)

        # ---- 页脚 ----
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=20, after=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"面试攻略 · {company_name} · {datetime.now().strftime('%Y-%m-%d')}")
        _set_run_font(run, "微软雅黑", 8, color=Colors.TEXT_LIGHT)

        # 保存
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        return str(output)

    def _style_table(self, table):
        """美化表格"""
        # 设置表格边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        # 交替行颜色
        for i, row in enumerate(table.rows):
            if i == 0:
                continue
            if i % 2 == 0:
                for cell in row.cells:
                    _set_cell_shading(cell, "F5F7FA")


# 全局实例
docx_writer = DocxWriter()
